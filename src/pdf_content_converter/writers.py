from __future__ import annotations

import html
import os
import re
import shutil
import subprocess
import tempfile
import uuid
import zipfile
from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image

from .models import OutputStyle, PageContent


_COMPOUND_FILE_SIGNATURE = bytes.fromhex("D0 CF 11 E0 A1 B1 1A E1")
_WORD_DOC_CONVERSION_SCRIPT = r'''param(
    [Parameter(Mandatory = $true)][string]$InputPath,
    [Parameter(Mandatory = $true)][string]$OutputPath
)

$ErrorActionPreference = 'Stop'
$word = $null
$document = $null
$originalAutomationSecurity = $null

try {
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $originalAutomationSecurity = $word.AutomationSecurity
    $word.AutomationSecurity = 3
    $document = $word.Documents.Open($InputPath, $false, $true, $false)
    $document.SaveAs2($OutputPath, 0, $false, '', $false)
    $document.Close(0)
    $document = $null
}
finally {
    if ($null -ne $document) {
        $document.Close(0)
    }
    if ($null -ne $word) {
        if ($null -ne $originalAutomationSecurity) {
            $word.AutomationSecurity = $originalAutomationSecurity
        }
        $word.Quit(0)
    }
    [GC]::Collect()
    [GC]::WaitForPendingFinalizers()
}

if (-not (Test-Path -LiteralPath $OutputPath)) {
    throw 'Microsoft Word did not create the requested DOC file.'
}
'''


def write_txt(pages: list[PageContent], destination: Path) -> Path:
    destination.parent.mkdir(parents=True, exist_ok=True)
    content = "\n\f\n".join(page.text for page in pages)
    _atomic_write_text(destination, content)
    return destination


def write_docx(
    pages: list[PageContent],
    destination: Path,
    title: str,
    style: OutputStyle,
) -> Path:
    destination.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.core_properties.title = title

    if style is OutputStyle.VISUAL:
        _write_visual_docx(document, pages)
    else:
        _write_editable_docx(document, pages)

    with tempfile.NamedTemporaryFile(
        prefix=f".{destination.stem}-",
        suffix=".docx",
        dir=destination.parent,
        delete=False,
    ) as temp_file:
        temp_path = Path(temp_file.name)
    try:
        document.save(temp_path)
        os.replace(temp_path, destination)
    finally:
        temp_path.unlink(missing_ok=True)
    return destination


def _write_editable_docx(document: Document, pages: list[PageContent]) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for page_index, page in enumerate(pages):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        lines = page.text.splitlines()
        if not lines:
            paragraph.add_run("")
        for line_index, line in enumerate(lines):
            if line_index:
                paragraph.add_run().add_break()
            paragraph.add_run(line)
        if page_index < len(pages) - 1:
            document.add_page_break()


def _write_visual_docx(document: Document, pages: list[PageContent]) -> None:
    for page_index, page in enumerate(pages):
        if page.image_path is None:
            raise ValueError(f"第 {page.number} 页缺少视觉保真图像。")
        section = (
            document.sections[0]
            if page_index == 0
            else document.add_section(WD_SECTION.NEW_PAGE)
        )
        page_width, page_height = _word_page_size(page.width_pt, page.height_pt)
        section.page_width = Inches(page_width)
        section.page_height = Inches(page_height)
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
        section.header_distance = Inches(0)
        section.footer_distance = Inches(0)

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(1)
        inline_shape = paragraph.add_run().add_picture(
            str(page.image_path),
            width=Inches(page_width),
            height=Inches(page_height),
        )
        _make_picture_page_anchored(inline_shape._inline)


def _word_page_size(width_pt: float, height_pt: float) -> tuple[float, float]:
    width = max(width_pt / 72.0, 1.0)
    height = max(height_pt / 72.0, 1.0)
    scale = min(1.0, 22.0 / width, 22.0 / height)
    return width * scale, height * scale


def _make_picture_page_anchored(inline) -> None:
    inline.tag = qn("wp:anchor")
    for name, value in {
        "distT": "0",
        "distB": "0",
        "distL": "0",
        "distR": "0",
        "simplePos": "0",
        "relativeHeight": "251658240",
        "behindDoc": "0",
        "locked": "1",
        "layoutInCell": "1",
        "allowOverlap": "1",
    }.items():
        inline.set(name, value)

    simple_position = OxmlElement("wp:simplePos")
    simple_position.set("x", "0")
    simple_position.set("y", "0")

    horizontal = OxmlElement("wp:positionH")
    horizontal.set("relativeFrom", "page")
    horizontal_offset = OxmlElement("wp:posOffset")
    horizontal_offset.text = "0"
    horizontal.append(horizontal_offset)

    vertical = OxmlElement("wp:positionV")
    vertical.set("relativeFrom", "page")
    vertical_offset = OxmlElement("wp:posOffset")
    vertical_offset.text = "0"
    vertical.append(vertical_offset)

    inline.insert(0, simple_position)
    inline.insert(1, horizontal)
    inline.insert(2, vertical)

    wrap_none = OxmlElement("wp:wrapNone")
    doc_properties = inline.find(qn("wp:docPr"))
    if doc_properties is None:
        raise ValueError("DOCX 图片锚点缺少 docPr。")
    inline.insert(inline.index(doc_properties), wrap_none)


def write_epub(
    pages: list[PageContent],
    destination: Path,
    title: str,
    style: OutputStyle,
) -> Path:
    destination.parent.mkdir(parents=True, exist_ok=True)
    book_id = f"urn:uuid:{uuid.uuid5(uuid.NAMESPACE_URL, title)}"
    page_files = [f"page-{page.number:05d}.xhtml" for page in pages]

    with tempfile.NamedTemporaryFile(
        prefix=f".{destination.stem}-",
        suffix=".epub",
        dir=destination.parent,
        delete=False,
    ) as temp_file:
        temp_path = Path(temp_file.name)

    try:
        with zipfile.ZipFile(temp_path, "w") as archive:
            archive.writestr(
                "mimetype",
                "application/epub+zip",
                compress_type=zipfile.ZIP_STORED,
            )
            archive.writestr(
                "META-INF/container.xml",
                _container_xml(),
                compress_type=zipfile.ZIP_DEFLATED,
            )

            image_items: list[tuple[str, str]] = []
            if style is OutputStyle.VISUAL:
                for page in pages:
                    if page.image_path is None:
                        raise ValueError(f"第 {page.number} 页缺少视觉保真图像。")
                    image_name = f"images/page-{page.number:05d}.png"
                    archive.write(page.image_path, f"OEBPS/{image_name}")
                    image_items.append((f"image-{page.number}", image_name))

            archive.writestr(
                "OEBPS/content.opf",
                _content_opf(book_id, title, page_files, image_items, style),
                compress_type=zipfile.ZIP_DEFLATED,
            )
            archive.writestr(
                "OEBPS/nav.xhtml",
                _nav_xhtml(title, page_files),
                compress_type=zipfile.ZIP_DEFLATED,
            )
            archive.writestr(
                "OEBPS/styles.css",
                _stylesheet(style),
                compress_type=zipfile.ZIP_DEFLATED,
            )
            for page, page_file in zip(pages, page_files, strict=True):
                archive.writestr(
                    f"OEBPS/{page_file}",
                    _page_xhtml(page, title, style),
                    compress_type=zipfile.ZIP_DEFLATED,
                )
        os.replace(temp_path, destination)
    finally:
        temp_path.unlink(missing_ok=True)
    return destination


def convert_docx_to_doc(docx_path: Path, destination: Path) -> Path:
    destination.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="pdf-doc-convert-") as temp_name:
        temp_dir = Path(temp_name)
        script_path = temp_dir / "convert-with-word.ps1"
        converted = temp_dir / f"{docx_path.stem}.doc"
        script_path.write_text(_WORD_DOC_CONVERSION_SCRIPT, encoding="utf-8-sig")
        completed = subprocess.run(
            [
                _find_windows_powershell(),
                "-NoProfile",
                "-NonInteractive",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                str(script_path),
                "-InputPath",
                str(docx_path.resolve()),
                "-OutputPath",
                str(converted.resolve()),
            ],
            capture_output=True,
            check=False,
            encoding="utf-8",
            errors="replace",
            timeout=180,
            creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )
        if completed.returncode != 0 or not converted.is_file():
            detail = (completed.stderr or completed.stdout).strip()
            raise RuntimeError(
                f"Microsoft Word 生成 DOC 失败: {detail or '没有输出文件'}"
            )
        with converted.open("rb") as stream:
            signature = stream.read(len(_COMPOUND_FILE_SIGNATURE))
        if signature != _COMPOUND_FILE_SIGNATURE:
            raise RuntimeError("Microsoft Word 未生成有效的 Word 97–2003 二进制 DOC。")
        shutil.copyfile(converted, destination)
    return destination


def _find_windows_powershell() -> str:
    candidates = [
        shutil.which("powershell.exe"),
        str(
            Path(os.environ.get("SystemRoot", r"C:\Windows"))
            / "System32"
            / "WindowsPowerShell"
            / "v1.0"
            / "powershell.exe"
        ),
    ]
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            return str(candidate)
    raise RuntimeError("找不到 Windows PowerShell，无法调用 Microsoft Word 生成 DOC。")


def _content_opf(
    book_id: str,
    title: str,
    page_files: list[str],
    image_items: list[tuple[str, str]],
    style: OutputStyle,
) -> str:
    modified = datetime.now(UTC).strftime("%Y-%m-%dT%H:%M:%SZ")
    manifest = [
        '<item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>',
        '<item id="style" href="styles.css" media-type="text/css"/>',
    ]
    spine = []
    for index, page_file in enumerate(page_files, start=1):
        manifest.append(
            f'<item id="page-{index}" href="{page_file}" media-type="application/xhtml+xml"/>'
        )
        spine.append(f'<itemref idref="page-{index}"/>')
    for item_id, href in image_items:
        manifest.append(
            f'<item id="{item_id}" href="{href}" media-type="image/png"/>'
        )
    rendition_metadata = ""
    if style is OutputStyle.VISUAL:
        rendition_metadata = (
            '<meta property="rendition:layout">pre-paginated</meta>'
            '<meta property="rendition:orientation">auto</meta>'
            '<meta property="rendition:spread">none</meta>'
        )
    return f'''<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="book-id">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:identifier id="book-id">{html.escape(book_id)}</dc:identifier>
    <dc:title>{html.escape(_xml_text(title))}</dc:title>
    <dc:language>und</dc:language>
    <meta property="dcterms:modified">{modified}</meta>
    {rendition_metadata}
  </metadata>
  <manifest>{''.join(manifest)}</manifest>
  <spine>{''.join(spine)}</spine>
</package>'''


def _page_xhtml(page: PageContent, title: str, style: OutputStyle) -> str:
    if style is OutputStyle.VISUAL:
        if page.image_path is None:
            raise ValueError(f"第 {page.number} 页缺少视觉保真图像。")
        image_width, image_height = _image_dimensions(page.image_path)
        viewport = (
            f'<meta name="viewport" content="width={image_width}, height={image_height}"/>'
        )
        body = (
            '<main class="fixed-page">'
            f'<img src="images/page-{page.number:05d}.png" '
            f'width="{image_width}" height="{image_height}" '
            f'alt="{html.escape(_xml_text(title))} - 第 {page.number} 页"/>'
            '</main>'
        )
    else:
        viewport = '<meta name="viewport" content="width=device-width, initial-scale=1.0"/>'
        safe_text = html.escape(_xml_text(page.text))
        body = f'<div class="page-text">{safe_text}</div>'
    return f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" lang="und">
  <head><title>{html.escape(_xml_text(title))}</title>{viewport}<link rel="stylesheet" type="text/css" href="styles.css"/></head>
  <body>{body}</body>
</html>'''


def _nav_xhtml(title: str, page_files: list[str]) -> str:
    links = "".join(
        f'<li><a href="{page_file}">第 {index} 页</a></li>'
        for index, page_file in enumerate(page_files, start=1)
    )
    return f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="zh-CN">
  <head><title>{html.escape(_xml_text(title))}</title></head>
  <body><nav epub:type="toc"><h1>{html.escape(_xml_text(title))}</h1><ol>{links}</ol></nav></body>
</html>'''


def _container_xml() -> str:
    return '''<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles><rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/></rootfiles>
</container>'''


def _stylesheet(style: OutputStyle) -> str:
    if style is OutputStyle.VISUAL:
        return (
            "html,body{width:100%;height:100%;margin:0;padding:0;overflow:hidden;}"
            ".fixed-page{position:absolute;inset:0;display:flex;align-items:center;"
            "justify-content:center;background:#fff;}"
            ".fixed-page img{display:block;width:100%;height:100%;object-fit:contain;}"
        )
    return (
        "body{font-family:serif;line-height:1.55;margin:5%;}"
        ".page-text{white-space:pre-wrap;overflow-wrap:anywhere;}"
    )


def _image_dimensions(path: Path) -> tuple[int, int]:
    with Image.open(path) as image:
        return image.size


def _xml_text(value: str) -> str:
    return re.sub(
        r"[^\x09\x0A\x0D\x20-\uD7FF\uE000-\uFFFD\U00010000-\U0010FFFF]",
        "",
        value,
    )


def _atomic_write_text(destination: Path, content: str) -> None:
    with tempfile.NamedTemporaryFile(
        prefix=f".{destination.stem}-",
        suffix=destination.suffix,
        dir=destination.parent,
        delete=False,
        mode="w",
        encoding="utf-8",
        newline="",
    ) as temp_file:
        temp_path = Path(temp_file.name)
        temp_file.write(content)
    try:
        os.replace(temp_path, destination)
    finally:
        temp_path.unlink(missing_ok=True)
