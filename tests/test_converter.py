from __future__ import annotations

import hashlib
import subprocess
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

import pypdfium2 as pdfium
from docx import Document
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

from pdf_content_converter.engine import _render_page, convert_pdf, discover_pdfs, find_word
from pdf_content_converter.gui import default_output_dir
from pdf_content_converter.models import ConversionOptions, OcrMode, OutputStyle
from pdf_content_converter.writers import convert_docx_to_doc


class ConverterTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls._workspace = tempfile.TemporaryDirectory(prefix="pdf-converter-tests-")
        cls.root = Path(cls._workspace.name)
        cls.source = cls.root / "sample.pdf"

        image_path = cls.root / "scan-page.png"
        image = Image.new("RGB", (1000, 1400), "white")
        draw = ImageDraw.Draw(image)
        draw.rectangle((60, 60, 940, 1340), outline="#1f4e78", width=12)
        draw.ellipse((180, 260, 820, 900), fill="#f4b183", outline="#c55a11", width=10)
        draw.line((140, 1150, 860, 980), fill="#70ad47", width=24)
        draw.text((80, 100), "SCANNED PAGE CONTENT 456", fill="black")
        image.save(image_path)
        image.close()

        pdf = canvas.Canvas(str(cls.source), pagesize=A4)
        pdf.drawString(72, 760, "Hello PDF content 123")
        pdf.setFillColorRGB(0.12, 0.31, 0.47)
        pdf.rect(72, 620, 450, 90, fill=1, stroke=0)
        pdf.setFillColorRGB(1, 1, 1)
        pdf.drawString(90, 675, "STRUCTURE + TABLE + IMAGE")
        pdf.setStrokeColorRGB(0.85, 0.85, 0.85)
        for x in (72, 222, 372, 522):
            pdf.line(x, 470, x, 590)
        for y in (470, 510, 550, 590):
            pdf.line(72, y, 522, y)
        pdf.drawImage(ImageReader(image_path), 180, 120, width=230, height=300)
        pdf.showPage()
        pdf.drawImage(str(image_path), 0, 0, width=A4[0], height=A4[1])
        pdf.save()

    @classmethod
    def tearDownClass(cls) -> None:
        cls._workspace.cleanup()

    def test_discover_pdfs_recursively(self) -> None:
        nested = self.root / "nested"
        nested.mkdir(exist_ok=True)
        nested_pdf = nested / "copy.PDF"
        nested_pdf.write_bytes(self.source.read_bytes())
        found = discover_pdfs([self.root], recursive=True)
        self.assertIn(self.source.resolve(), found)
        self.assertIn(nested_pdf.resolve(), found)

    def test_text_outputs_and_source_are_unchanged(self) -> None:
        before = _sha256(self.source)
        output = self.root / "text-output"
        options = ConversionOptions(
            output_dir=output,
            formats=("txt", "epub", "docx"),
            ocr_mode=OcrMode.NEVER,
            output_style=OutputStyle.EDITABLE,
        )
        result = convert_pdf(self.source, options)

        self.assertEqual({}, result.errors)
        self.assertEqual(before, _sha256(self.source))
        self.assertEqual({"txt", "epub", "docx"}, set(result.outputs))
        self.assertIn("Hello PDF content 123", result.outputs["txt"].read_text("utf-8"))

        document = Document(result.outputs["docx"])
        docx_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Hello PDF content 123", docx_text)

        with zipfile.ZipFile(result.outputs["epub"]) as archive:
            self.assertEqual("mimetype", archive.namelist()[0])
            self.assertEqual(
                b"application/epub+zip",
                archive.read("mimetype"),
            )
            self.assertIn(
                "Hello PDF content 123",
                archive.read("OEBPS/page-00001.xhtml").decode("utf-8"),
            )

    def test_visual_outputs_contain_both_pages(self) -> None:
        output = self.root / "visual-output"
        options = ConversionOptions(
            output_dir=output,
            formats=("epub", "docx"),
            ocr_mode=OcrMode.AUTO,
            output_style=OutputStyle.VISUAL,
            dpi=120,
        )
        with patch("pdf_content_converter.engine.find_tesseract", return_value=None):
            result = convert_pdf(self.source, options)
        self.assertEqual({}, result.errors)

        Document(result.outputs["docx"])
        with zipfile.ZipFile(result.outputs["docx"]) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
            self.assertEqual(2, document_xml.count("<wp:anchor"))
            docx_images = sorted(
                archive.read(name)
                for name in archive.namelist()
                if name.startswith("word/media/")
            )
            self.assertEqual(2, len(docx_images))
        with zipfile.ZipFile(result.outputs["epub"]) as archive:
            images = [name for name in archive.namelist() if name.startswith("OEBPS/images/")]
            self.assertEqual(2, len(images))
            package = archive.read("OEBPS/content.opf").decode("utf-8")
            self.assertIn('<meta property="rendition:layout">pre-paginated</meta>', package)
            self.assertIn('<meta property="rendition:spread">none</meta>', package)
            page = archive.read("OEBPS/page-00001.xhtml").decode("utf-8")
            self.assertIn('name="viewport"', page)
            epub_images = [archive.read(name) for name in sorted(images)]

        baseline_dir = self.root / "visual-baseline"
        baseline_dir.mkdir(exist_ok=True)
        document = pdfium.PdfDocument(str(self.source))
        try:
            baseline_paths = []
            for index in range(len(document)):
                path = baseline_dir / f"page-{index + 1:05d}.png"
                _render_page(document, index, path, 120)
                baseline_paths.append(path)
        finally:
            document.close()
        baseline_images = [path.read_bytes() for path in baseline_paths]
        self.assertEqual(
            [_sha256_bytes(data) for data in baseline_images],
            [_sha256_bytes(data) for data in epub_images],
        )
        self.assertEqual(
            sorted(_sha256_bytes(data) for data in baseline_images),
            sorted(_sha256_bytes(data) for data in docx_images),
        )

    def test_fidelity_is_the_default_style(self) -> None:
        options = ConversionOptions(output_dir=self.root / "default-output")
        self.assertIs(OutputStyle.VISUAL, options.output_style)

    def test_gui_default_output_is_desktop(self) -> None:
        self.assertEqual(Path.home() / "Desktop", default_output_dir())

    def test_auto_ocr_fails_clearly_without_tesseract(self) -> None:
        options = ConversionOptions(
            output_dir=self.root / "missing-ocr-output",
            formats=("txt",),
            ocr_mode=OcrMode.AUTO,
        )
        with patch("pdf_content_converter.engine.find_tesseract", return_value=None):
            result = convert_pdf(self.source, options)
        self.assertIn("conversion", result.errors)
        self.assertIn("Tesseract", result.errors["conversion"])
        self.assertFalse(result.outputs)

    def test_full_ocr_processes_every_rendered_page(self) -> None:
        options = ConversionOptions(
            output_dir=self.root / "full-ocr-output",
            formats=("txt",),
            ocr_mode=OcrMode.FULL,
            dpi=120,
        )
        with (
            patch("pdf_content_converter.engine.find_tesseract", return_value="tesseract"),
            patch(
                "pdf_content_converter.engine._run_tesseract",
                side_effect=["OCR PAGE ONE", "OCR PAGE TWO"],
            ),
        ):
            result = convert_pdf(self.source, options)
        self.assertEqual({}, result.errors)
        self.assertEqual([1, 2], result.ocr_pages)
        text = result.outputs["txt"].read_text("utf-8")
        self.assertIn("OCR PAGE ONE", text)
        self.assertIn("OCR PAGE TWO", text)

    def test_missing_microsoft_word_only_blocks_doc(self) -> None:
        options = ConversionOptions(
            output_dir=self.root / "missing-doc-output",
            formats=("txt", "doc"),
            ocr_mode=OcrMode.NEVER,
        )
        with patch("pdf_content_converter.engine.find_word", return_value=None):
            result = convert_pdf(self.source, options)
        self.assertIn("txt", result.outputs)
        self.assertIn("doc", result.errors)
        self.assertIn("Microsoft Word", result.errors["doc"])
        self.assertNotIn("doc", result.outputs)

    def test_find_word_accepts_an_explicit_executable(self) -> None:
        fake_word = self.root / "WINWORD.EXE"
        fake_word.touch()
        self.assertEqual(str(fake_word), find_word(str(fake_word)))

    def test_doc_conversion_uses_word_automation_and_binary_format(self) -> None:
        source_docx = self.root / "word-input.docx"
        destination = self.root / "word-output.doc"
        document = Document()
        document.add_paragraph("Microsoft Word conversion test")
        document.save(source_docx)
        captured: dict[str, str] = {}

        def fake_run(command, **kwargs):
            script_path = Path(command[command.index("-File") + 1])
            output_path = Path(command[command.index("-OutputPath") + 1])
            captured["script"] = script_path.read_text(encoding="utf-8-sig")
            output_path.write_bytes(bytes.fromhex("D0CF11E0A1B11AE1") + b"test-doc")
            return subprocess.CompletedProcess(command, 0, "", "")

        with (
            patch(
                "pdf_content_converter.writers._find_windows_powershell",
                return_value="powershell.exe",
            ),
            patch("pdf_content_converter.writers.subprocess.run", side_effect=fake_run),
        ):
            result = convert_docx_to_doc(source_docx, destination)

        self.assertEqual(destination, result)
        self.assertTrue(destination.read_bytes().startswith(bytes.fromhex("D0CF11E0A1B11AE1")))
        self.assertIn("Word.Application", captured["script"])
        self.assertIn("AutomationSecurity = 3", captured["script"])
        self.assertIn("SaveAs2($OutputPath, 0", captured["script"])

    def test_doc_conversion_rejects_non_binary_output(self) -> None:
        source_docx = self.root / "invalid-word-input.docx"
        destination = self.root / "invalid-word-output.doc"
        Document().save(source_docx)

        def fake_run(command, **kwargs):
            output_path = Path(command[command.index("-OutputPath") + 1])
            output_path.write_text("not a real DOC", encoding="utf-8")
            return subprocess.CompletedProcess(command, 0, "", "")

        with (
            patch(
                "pdf_content_converter.writers._find_windows_powershell",
                return_value="powershell.exe",
            ),
            patch("pdf_content_converter.writers.subprocess.run", side_effect=fake_run),
        ):
            with self.assertRaisesRegex(RuntimeError, "Word 97–2003"):
                convert_docx_to_doc(source_docx, destination)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


if __name__ == "__main__":
    unittest.main()
